import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
import string
from collections import Counter
import re
import subprocess
import spacy


nlp = spacy.load(r"C:\Users\Alex\AppData\Local\Programs\Python\Python310\Lib\site-packages\pt_core_news_sm\pt_core_news_sm-3.7.0")

nltk.download('stopwords')
nltk.download('punkt')

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def preprocess_text(text):
    # Tokenize o texto em palavras
    words = word_tokenize(text, language='portuguese')

    # Remova stop words, dígitos e pontuação
    filtered_words = [word.lower() for word in words if
                      word.lower() not in stop_words and word not in string.punctuation and not word.isdigit()]

    # Recrie o texto a partir das palavras filtradas
    filtered_text = ' '.join(filtered_words)

    return filtered_text


def is_pdf_scanned(pdf_path):
    pdf_reader = PdfReader(pdf_path)

    for page in pdf_reader.pages:
        text = page.extract_text()
        if text.strip():
            return False
    return True


def extract_text_from_image(image_path):
    return pytesseract.image_to_string(Image.open(image_path))


def detect_paragraphs(text):
    return text.split('\n\n')


def remove_camscanner_line(text):
    return text.replace("Digitalizado com CamScanner", "")


pdf_path = 'pdf.pdf'


def format_text(text, paragraph_margin='    '):
    paragraphs = text.split('\n\n')
    formatted_text = ""

    for paragraph in paragraphs:
        if "Digitalizado com CamScanner" in paragraph:
            continue

        lines = paragraph.split('\n')
        formatted_paragraph = ""

        for line in lines:
            # Remover espaços em branco no início
            line = line.strip()

            words = line.split()
            if words:
                current_line = words[0]

                for word in words[1:]:
                    if word.startswith('-'):
                        # Juntar palavras divididas com hífen
                        current_line = current_line.rstrip('-') + word
                    else:
                        current_line += ' ' + word

                formatted_paragraph += current_line + ' '

        # Adicionar margem no início de cada parágrafo
        formatted_paragraph = paragraph_margin + formatted_paragraph

        formatted_text += formatted_paragraph + '\n'

    return formatted_text


if is_pdf_scanned(pdf_path):
    print("Usando pytesseract para PDFs escaneados")

    pdf_pages = convert_from_path(pdf_path, 500)
    all_text = ""

    for page_enumetarion, page in enumerate(pdf_pages, start=1):
        filename = f"page_{page_enumetarion:03}.jpg"
        page.save(filename, "JPEG")
        text = extract_text_from_image(filename)
        all_text += text

    all_text = remove_camscanner_line(all_text)

    formatted_text = format_text(all_text)

    doc = Document()
    doc.add_paragraph(formatted_text)

    doc.save('saida.docx')
else:
    print("Usando PyPDF2 para PDFs nativos")

    pdf_reader = PdfReader(pdf_path)
    all_text = ""

    for page_number, page in enumerate(pdf_reader.pages, start=1):
        text = page.extract_text()
        all_text += text

    all_text = remove_camscanner_line(all_text)

    formatted_text = format_text(all_text)

    doc = Document()
    doc.add_paragraph(formatted_text)

    doc.save('saida.docx')

doc = Document('saida.docx')
stop_words = set(stopwords.words('portuguese'))
processed_paragraphs = []

for paragraph in doc.paragraphs:

    cleaned_paragraph = paragraph.text.replace('\n', '').strip()

    processed_text = preprocess_text(cleaned_paragraph)

    # Usar spaCy para análise de partes do discurso e eliminar pronomes
    doc = nlp(processed_text)
    processed_text = ' '.join([token.text for token in doc if token.pos_ != 'PRON'])

    processed_paragraphs.append(processed_text)

new_doc = Document()
for processed_paragraph in processed_paragraphs:
    new_doc.add_paragraph(processed_paragraph)

# Salvar o novo documento em um arquivo
new_doc.save('saida_processado.docx')

doc = Document('saida_processado.docx')
all_words = []

for paragraph in doc.paragraphs:
    text = paragraph.text

    words = re.findall(r'\b\w+\b', text.lower())
    all_words.extend(words)

word_count = Counter(all_words)

most_common_words = word_count.most_common(10)

for word, count in most_common_words:
    print(f'{word}: {count}')

subprocess.run(["python", "tradutor.py", "text_to_speech.py"])